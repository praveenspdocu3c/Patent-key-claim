import streamlit as st
from openai import AzureOpenAI
from azure.ai.formrecognizer import DocumentAnalysisClient # type: ignore
from azure.core.credentials import AzureKeyCredential
from docx import Document
import os

# hide_streamlit_style = """
#             <style>
#             [data-testid="stToolbar"] {visibility: hidden !important;}
#             footer {visibility: hidden !important;}
#             </style>
#             """
# st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# Inject CSS to hide the avatar
hide_avatar_style = """
    <style>
    ._profileImage_51w34_76 {
        display: none !important;
    }
    </style>
"""
st.markdown(hide_avatar_style, unsafe_allow_html=True)

# hide_profile_js = """
#     <script>
#     // Hides the profile icon button
#     const profileIcon = document.querySelector('iframe[title="profile"]');
#     if (profileIcon) {
#         profileIcon.style.display = 'none';
#     }
#     </script>
#     """
# # Inject the JS into the Streamlit app
# st.markdown(hide_profile_js, unsafe_allow_html=True)

# hide_streamlit_style = """
#             <style>
#             /* Hide the header toolbar */
#             [data-testid="stToolbar"] {visibility: hidden !important;}
            
#             /* Hide Streamlit's profile icon and the footer */
#             footer {visibility: hidden !important;}
#             footer:after {content:''; visibility: hidden;}
            
#             /* Attempt to hide just the profile icon */
#             [data-testid="stDecoration"] {visibility: hidden !important;}
#             </style>
#             """
# st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# Azure OpenAI setup
client = AzureOpenAI(
    azure_endpoint="https://theswedes.openai.azure.com/",
    api_key="783973291a7c4a74a1120133309860c0",
    api_version="2024-02-01",
)

# Function to use Azure Document Intelligence for PDF extraction
def extract_text_from_pdf(uploaded_pdf):
    endpoint = "https://patentocr.cognitiveservices.azure.com/"
    api_key = "cd6b8996d93447be88d995729c924bcb"

    document_analysis_client = DocumentAnalysisClient(
        endpoint=endpoint,
        credential=AzureKeyCredential(api_key),
    )

    poller = document_analysis_client.begin_analyze_document(
        "prebuilt-document", document=uploaded_pdf
    )
    result = poller.result()

    text = ""
    for page in result.pages:
        for line in page.lines:
            text += line.content + "\n"
    
    return text

def check_for_conflicts(action_document_text):
    prompt = f"""
    Analyze the following action document text and check for conflicts under U.S.C. 102 and 103:

    {action_document_text}

    Step 1: Extract the Key Claim(Independent Claims) ((Extract with the entire technical content)) and consider it as 'Key_claims'.
    Step 2: Extract the Root Claim(Independent Claims) presented in the 'Key_claims' and map them accordingly. (Note: Please note method claims and system claims are not consider as independent Claims)
    Step 3: Extract all referenced documents under U.S.C. 102 and 103 mentioned in the action document based on "Root Claim(Independent Claims)" is presented.
    Step 4: For each referenced document, create a variable that stores the document name.
    Step 5: If any Root Claim(Independent Claims) is being anticipated by the referenced documents, extract the entire technical content with its specified paragraph location and image reference. Then, map the claim with the conflicting document name.
    Return the Root Claim(Independent Claims) and their mapping to conflicting documents.
    """

    messages = [
        {
            "role": "system",
            "content": """  
            You are a patent attorney.
            Analyze the following action document text and check for conflicts under Section 102 and 103. Follow the steps to extract referenced documents and their associated Key Claims, then map them to conflicting documents.
            """,
        },
        {
            "role": "user",
            "content": prompt,
        },
    ]

    response = client.chat.completions.create(
        model="GPT-4-Omni", messages=messages, temperature=0.6
    )

    output = response.choices[0].message.content
    print("_________________________________________________________________________________________________________")
    print(output)
    print("_________________________________________________________________________________________________________")


    if "anticipated by" in output:
        claim = "Claim 1"
        documents = [doc.strip() for doc in output.split("anticipated by")[1].split(",")]
        return True, claim, documents

    return False, "", []

# Function to compare Claim 1 against other documents
# Function to compare Claim 1 and Original Document against other documents

def compare_claims(claim1_text, other_docs_texts):
    results = {}
    
    for doc_name, doc_text in other_docs_texts.items():
        messages = [
            {
                "role": "system",
                "content": """You are an expert patent attorney, analyzing patent claims based on Claim Rejections sections 102 and 103 of the US Patent Act.
                              Determine if the examiner is correct in asserting that Claim 1 is anticipated by the cited reference of 'Cited Reference' comparing with Application as filed as U.S pre-grant Publication document 'Claim 1'.""",
            },
            {
                "role": "user",
                "content": f"""
                Compare Claim 1 with the cited reference. 
                \n\nClaim 1: {claim1_text}
                \n\nCited Reference: {doc_text}
                """,
            },
        ]

        response = client.chat.completions.create(
            model="GPT-4-Omni", messages=messages, temperature=0 
        )
        
        # Capture the entire response
        full_output = response.choices[0].message.content
        results[doc_name] = full_output
    
    return results

from docx import Document

# Function to generate Word document with results, converting **text** into bold
def generate_word_doc(comparison_results):
    doc = Document()
    doc.add_heading("Patent Comparison Results", 0)
    
    for doc_name, result in comparison_results.items():
        # Add document name in bold instead of using a heading
        p = doc.add_paragraph()
        p.add_run(f"Analysis for {doc_name}").bold = True
        
        # Process the result content and convert **text** into bold text
        lines = result.split("\n")
        for line in lines:
            p = doc.add_paragraph()  # Create a new paragraph for each line
            words = line.split("**")
            for i, word in enumerate(words):
                if i % 2 == 0:
                    # Add normal text
                    p.add_run(word)
                else:
                    # Add bold text
                    p.add_run(word).bold = True
    
    doc_filename = "comparison_results.docx"
    doc.save(doc_filename)
    return doc_filename

# Streamlit App UI
st.title("Patent Key Claims Analysis")

# Initialize session state to store the uploaded files
if "action_document_pdf" not in st.session_state:
    st.session_state.action_document_pdf = None
if "original_document_pdf" not in st.session_state:
    st.session_state.original_document_pdf = None
if "other_documents_pdfs" not in st.session_state:
    st.session_state.other_documents_pdfs = []

if "conflict_detected" not in st.session_state:
    st.session_state.conflict_detected = False
if "claim_text" not in st.session_state:
    st.session_state.claim_text = None
if "conflict_docs" not in st.session_state:
    st.session_state.conflict_docs = []

# Upload the action document
st.sidebar.title("Upload Action Document (NFOA)")
uploaded_action_document = st.sidebar.file_uploader("Upload Action Document (NFOA) PDF", type=["pdf"])

# Analyze Action Document for Conflicts
if uploaded_action_document is not None and st.sidebar.button("Analyze Action Document"):
    st.session_state.action_document_pdf = uploaded_action_document
    action_document_text = extract_text_from_pdf(st.session_state.action_document_pdf)
    conflict_exists, claim_text, conflict_docs = check_for_conflicts(action_document_text)

    if conflict_exists:
        st.session_state.conflict_detected = True
        st.session_state.claim_text = claim_text
        st.session_state.conflict_docs = conflict_docs
        st.success(f"Key Claim detected for {uploaded_action_document.name}. Claim rejection documents: {', '.join(conflict_docs)}")
    else:
        st.error(f"Please try again, No Key Claim detected for {uploaded_action_document.name}")

# After detecting conflict, ask for the original and conflicting documents
if st.session_state.conflict_detected:
    st.sidebar.title("Upload Documents for Comparison")
    
    # Upload Original Document
    original_document = st.sidebar.file_uploader("Upload the Application as Filed Document", type=["pdf"], key="original")

    if original_document is not None:
        st.session_state.original_document_pdf = original_document
    
    # Upload Other Documents (Batch Mode)
    other_documents = st.sidebar.file_uploader("Claim Rejection Analysis Patent Documents", type=["pdf"], key="batch", accept_multiple_files=True)

    if other_documents:
        st.session_state.other_documents_pdfs = other_documents

    # Once documents are uploaded, enable comparison
    if st.session_state.original_document_pdf and st.session_state.other_documents_pdfs and st.sidebar.button("Compare Documents"):
        original_text = extract_text_from_pdf(st.session_state.original_document_pdf)
        other_docs_texts = {f"Claim rejection analysis {i+1}. {doc.name}": extract_text_from_pdf(doc) for i, doc in enumerate(st.session_state.other_documents_pdfs)}

        # Compare Claim 1 and Original Document against the other documents
        comparison_results = compare_claims(st.session_state.claim_text, other_docs_texts)

        # Generate the Word document with the comparison results
        doc_filename = generate_word_doc(comparison_results)

        # Display the comparison results
        # for doc_name, results in comparison_results.items():
        #     st.subheader(f"Analysis with {doc_name}")
        #     st.write(f"**Claim 1 Comparison**: {results['Claim 1 Comparison']}")
        #     st.write(f"**Original Document Comparison**: {results['Original Document Comparison']}")
            
        for doc_name, result in comparison_results.items():
            st.subheader(f"Analysis with {doc_name}")
            st.write(result)


        # Provide download link for Word document
        with open(doc_filename, "rb") as f:
            st.download_button(
                label="Download Comparison Results",
                data=f,
                file_name=doc_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
